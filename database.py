import pandas as pd
import sqlite3
from io import BytesIO
from docx import Document
import os
import fitz



class database:
    def __init__(self, db_name):
        self.db_name = db_name
        self.connection = None
        self.cursor = None

    def connect(self):
        self.connection = sqlite3.connect(self.db_name)
        self.cursor = self.connection.cursor()

    def disconnect(self):
        if self.connection:
            self.connection.close()
            self.connection = None

    def is_connected(self):
        return self.connection is not None
    


    def readQuestionnaire(self):
        selected_questionnaire = "SELECT * FROM Questionnaire"
        return pd.read_sql_query(selected_questionnaire, self.connection)
    
    def readQuestionnaireClean(self):
        selected_questionnaire = "SELECT * FROM QuestionnaireClean"
        return pd.read_sql_query(selected_questionnaire, self.connection)
    def execute_query(self, query, params=None):
        if not self.is_connected():
            raise Exception("Database connection is not established.")
        
        cursor = self.connection.cursor()
        if params:
            cursor.execute(query, params)
        else:
            cursor.execute(query)
        
        self.connection.commit()
        return cursor.fetchone()
    


    def retrieve_and_save_doc(self,filename):

        record = self.execute_query( 'SELECT file_data FROM cahierDeCharges WHERE filename = ?', (filename,))
        
        if record:
            file_data = record[0]
            return file_data
        else:
            return None
    
    def readCahierDeCharges(self, filename):
        CahierCharge = self.retrieve_and_save_doc(filename)

        if CahierCharge is not None:
            document = Document(BytesIO(CahierCharge))
            full_text = "\n".join(p.text.strip() for p in document.paragraphs if p.text.strip())
            return full_text


        local_path = f'./clients/{filename}'

        if os.path.exists(local_path):
            if filename.lower().endswith(".docx"):
                document = Document(local_path)
                full_text = "\n".join(p.text.strip() for p in document.paragraphs if p.text.strip())
                return full_text
            elif filename.lower().endswith(".pdf"):

               
                with fitz.open(local_path) as doc:
                    full_text = "\n".join(page.get_text().strip() for page in doc)
                return full_text
            else:
                print(f"❌ Format de fichier non supporté : {filename}")
                return ""
        else:
            print(f"⚠️ Fichier '{filename}' non trouvé dans la base de données ni dans ./cahiercharge.")
            return ""


    def read_json(self, filename):
        """
        Lit un fichier JSON depuis la base de données et le retourne sous forme de DataFrame.
        """
        record = self.execute_query('SELECT json_content FROM ResultatsJSON WHERE filename = ?', (filename,))
        
        if record:
            json_content = record[0]
            from io import StringIO  

            return pd.read_json(StringIO(json_content))
        else:
            print(f"⚠️ Fichier '{filename}' non trouvé dans la base de données.")
            return pd.DataFrame()
        
    def commit(self):
        """
        Commit les changements dans la base de données.
        """
        if self.connection:
            self.connection.commit()
        else:
            raise Exception("Database connection is not established.")
    
    def close(self):
        """
        Ferme la connexion à la base de données.
        """
        if self.connection:
            self.connection.close()
            self.connection = None
        else:
            raise Exception("Database connection is not established.")